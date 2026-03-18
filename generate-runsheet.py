#!/usr/bin/env python3
"""Generate the Claude 101 Session Run Sheet as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# Kaizen colors
INK = RGBColor(0x1A, 0x1A, 0x18)
GOLD = RGBColor(0xC9, 0xA8, 0x4C)
STONE = RGBColor(0x8C, 0x87, 0x78)
CHARCOAL = RGBColor(0x2E, 0x2E, 0x2A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = INK

# Adjust margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def add_title(text, size=28):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = INK
    run.bold = True
    p.space_after = Pt(4)
    return p

def add_subtitle(text, size=14):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = STONE
    p.space_after = Pt(16)
    return p

def add_section_header(text, size=18):
    p = doc.add_paragraph()
    p.space_before = Pt(24)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = INK
    run.bold = True
    p.space_after = Pt(4)
    # Gold underline bar
    bar = doc.add_paragraph()
    run2 = bar.add_run('━' * 8)
    run2.font.color.rgb = GOLD
    run2.font.size = Pt(10)
    bar.space_after = Pt(12)
    return p

def add_phase_header(phase_num, title, time_range, size=16):
    p = doc.add_paragraph()
    p.space_before = Pt(32)
    # Phase label
    run = p.add_run(f'PHASE {phase_num}: ')
    run.font.size = Pt(10)
    run.font.color.rgb = GOLD
    run.bold = True
    run.font.name = 'Calibri'
    p.add_run('\n')
    # Title
    run2 = p.add_run(title)
    run2.font.size = Pt(size)
    run2.font.color.rgb = INK
    run2.bold = True
    # Time
    run3 = p.add_run(f'  ({time_range})')
    run3.font.size = Pt(11)
    run3.font.color.rgb = STONE
    p.space_after = Pt(8)
    # Divider
    bar = doc.add_paragraph()
    run4 = bar.add_run('━' * 40)
    run4.font.color.rgb = GOLD
    run4.font.size = Pt(8)
    bar.space_after = Pt(12)
    return p

def add_body(text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    p.space_after = Pt(8)
    return p

def add_script(text):
    """Add a quoted script block"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = CHARCOAL
    run.italic = True
    return p

def add_guardrail(text):
    """Add a guardrail callout"""
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    run_label = p.add_run('GUARDRAIL: ')
    run_label.font.size = Pt(9)
    run_label.font.color.rgb = GOLD
    run_label.bold = True
    run_text = p.add_run(text)
    run_text.font.size = Pt(10)
    run_text.font.color.rgb = STONE
    p.space_after = Pt(12)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(11)
        run_t = p.add_run(text)
        run_t.font.size = Pt(11)
    else:
        p.clear()
        run_t = p.add_run(text)
        run_t.font.size = Pt(11)
    p.space_after = Pt(4)
    return p

def add_checkbox(text):
    p = doc.add_paragraph()
    run = p.add_run('☐  ' + text)
    run.font.size = Pt(11)
    p.space_after = Pt(4)
    return p

# ═══════════════════════════════════════════
# DOCUMENT START
# ═══════════════════════════════════════════

add_title('Claude 101')
add_subtitle('Session Run Sheet — 90-Minute Workshop')

p = doc.add_paragraph()
run = p.add_run('Deliverable: ')
run.bold = True
run.font.size = Pt(12)
run2 = p.add_run('Claude account + Business Profile document uploaded to a Project')
run2.font.size = Pt(12)
p.space_after = Pt(8)

p = doc.add_paragraph()
run = p.add_run('Guide URL: ')
run.bold = True
run = p.add_run('claude.kaizencollective.com.au')
run.font.color.rgb = GOLD
p.space_after = Pt(24)

# ─── PRE-SESSION CHECKLIST ───
add_section_header('Pre-Session Checklist')
add_checkbox('Screen share ready — claude.kaizencollective.com.au open in browser')
add_checkbox('Timer app ready (visible countdown for build sprints)')
add_checkbox('Send message 24hrs before: "Tomorrow\'s session is hands-on. You\'ll need a laptop and you\'ll be signing up for Claude. Come ready to build."')
add_checkbox('Know 3-4 client names to cold-call during the session')
add_checkbox('Have your own Claude account open for live demos')
add_checkbox('Homework text drafted and ready to paste into Slack after the call')

# ─── PHASE 1: OPEN ───
add_phase_header(1, 'OPEN', '0:00–0:05 — 5 min')

add_body('Wins Round (3 min)')
add_script('"Chuck into the chat: what\'s your biggest win since we last spoke?"')
add_body('Call on 2-3 people by name. Celebrate genuinely. Keep it moving.')

add_body('Session Roadmap (2 min)')
add_body('Screen share the guide — show the sidebar nav.')
add_script('"Here\'s what we\'re building today. Three things:\n1. You\'ll understand why Claude is different from ChatGPT — and feel it\n2. You\'ll set up your Claude account and try it live on the call\n3. You\'ll walk away with your first Business Profile document\n\nBy the end of this call, you\'ll have a thinking partner that knows your business better than most people on your team."')

add_guardrail('No banter before the wins round. 5 minutes total. Hard cap.')

# ─── PHASE 2: FRAME ───
add_phase_header(2, 'FRAME', '0:05–0:15 — 10 min')

add_body('Paint the Pain (4 min)')
add_body('Scroll to the Introduction page hero. Read/riff on the opening copy, then go off-script with your lived experience:')
add_script('"I was the same. I had ChatGPT open in a tab. I\'d ask it to write an email. It was... fine. But I knew I was using 5% of what was possible. Then I found Claude. And in two weeks I built 11 applications, 16 documents, and compressed 6 months of development work into days. Without writing code."')
add_body('Show the stat strip: 11 apps. 16+ documents. 6 months compressed. 0 lines of code.')

add_body('Name the Pattern (2 min)')
add_script('"Most of you are using AI like a fancy search engine. Type a question, get an answer, move on. That\'s not wrong — it\'s just the microwave setting. There\'s a commercial kitchen here and you haven\'t turned on the stove."')

add_body('Bridge — Transition Question (2 min)')
add_script('"Be honest — chuck a number in the chat. On a scale of 1-10, how much of AI\'s potential do you think you\'re actually using right now?"')
add_body('Wait for responses. React to a few.')
add_script('"Right. Most of you are at a 2 or 3. By the end of today you\'ll be at a 6. By the end of the month, you\'ll be at an 8. Let me show you why."')

add_guardrail('Don\'t spend more than 10 minutes here. The temptation is to keep talking about how good Claude is. Resist. Let them FEEL it instead.')

# ─── PHASE 3: TEACH ───
add_phase_header(3, 'TEACH — The Comparison', '0:15–0:25 — 10 min')

add_body('Show the Destination First (2 min)')
add_body('Click to "What I\'ve Built" tab. Scroll through the showcase:')
add_bullet('The Lighthouse screenshot (pause here — let it sink in)')
add_bullet('CRM Dashboard, Onboarding Deck, Daily Slack Report')
add_bullet('The two landing pages (click through to artifacts)')
add_script('"Everything you\'re looking at was built by me and Claude. I\'m not a developer. I described what I wanted in plain English."')

add_body('The Prompt Showdown (5 min)')
add_body('Click to "ChatGPT vs Claude" tab. Screen share the three scenarios:')
add_body('1. Offer Strategy — read both outputs aloud. Pause on Claude\'s diagnostic response.')
add_script('"See that? ChatGPT gave a template. Claude asked what the actual goal was."')
add_body('2. Staff Problem — quick scroll, highlight the verdict.')
add_body('3. Content Strategy — this one will resonate most.')
add_script('"\'Post more\' isn\'t a strategy. What\'s actually broken?"')

add_body('Cold-Call Recall (2 min)')
add_script('"[Name], what did Claude do differently in all three of those examples?"')
add_body('Wait for answer. Reinforce: "It asked questions first. It diagnosed before it prescribed."')

add_guardrail('Don\'t read every word on the page. Hit the highlights. The page is their reference for later.')

# ─── PHASE 4: BUILD #1 ───
add_phase_header(4, 'BUILD SPRINT #1 — Sign Up + Try Raw', '0:25–0:45 — 20 min')

add_body('Setup Instruction (2 min)')
add_body('Click to "Getting Started" tab.')
add_script('"Alright. Laptops open. Go to claude.ai right now. Sign up with your business email. Free plan is fine for today. You have 3 minutes. Type DONE in the chat when you\'re in."')
add_body('⏱ START TIMER: 3 minutes')
add_body('Walk the room virtually. Help anyone stuck on signup.')

add_body('Try It Now — The Three Prompts (15 min)')
add_body('Click to "Try It Now" tab.')
add_script('"Now here\'s where it gets fun. Claude has NO context about your business right now. It\'s working blind. Copy each prompt, paste it in, and see what happens."')

add_body('Prompt 1: The Question Test (5 min)')
add_script('"Copy the first prompt. Replace the brackets with your real business and numbers. Paste it into Claude. Watch what happens."')
add_body('⏱ START TIMER: 5 minutes')
add_body('Walk the room. React: "[Name], what did Claude ask you? See? It didn\'t give you a list. It asked for context first."')

add_body('Prompt 2: The Voice Test (5 min)')
add_script('"Find something you\'ve written recently — a caption, an email, whatever. Paste it in with the second prompt."')
add_body('⏱ START TIMER: 5 minutes')

add_body('Prompt 3: The Pushback Test (5 min)')
add_script('"Think of a decision you\'re wrestling with right now. A real one. Describe it and ask Claude to find the holes."')
add_body('⏱ START TIMER: 5 minutes')

add_body('Show-and-Tell (3 min)')
add_script('"Who got something that surprised them? [Name], share your screen."')
add_body('Pick 2 people max.')
add_script('"That was Claude working BLIND. No context. Imagine what happens when we load your full business profile, your brand voice, your client data. That\'s what we\'re doing next."')

add_guardrail('Max 3 walkthroughs. If everyone wants to share, tell them to screenshot and post in Slack.')

# ─── PHASE 5: BUILD #2 ───
add_phase_header(5, 'BUILD SPRINT #2 — Business Profile', '0:45–1:10 — 25 min')

add_body('Frame the Prompt Vault (3 min)')
add_body('Click to "The Prompt Vault" tab.')
add_script('"These five prompts are engineered. Each one forces Claude into diagnostic mode — it asks you questions one at a time, goes deep, and produces a structured document at the end. We\'re starting with the Business Profile. This is the foundation."')

add_body('Build Sprint (20 min)')
add_script('"Copy Prompt 1 — the Business Profile Builder. Paste it into Claude. Answer every question Claude asks you. Be specific — if Claude asks about your business model, don\'t say \'I do fitness.\' Tell it your actual revenue, your actual structure, your actual capacity. You have 20 minutes. Type DONE when Claude produces your finished document."')
add_body('⏱ START TIMER: 20 minutes')

add_body('During the sprint:')
add_bullet('Check in by name every 3-4 minutes: "[Name], how\'s it going? What question are you up to?"')
add_bullet('If someone finishes early: "Start on Prompt 2 — the ICP."')
add_bullet('If someone is stuck: "Read me Claude\'s last question. What are you unsure about?"')
add_bullet('Celebrate progress: "How good is that? Claude just mapped your entire business in 15 minutes."')

add_body('Show-and-Tell (2 min)')
add_script('"Who\'s got a finished Business Profile? [Name], share your screen."')
add_body('Show one example. React to the specificity.')
add_script('"That took 20 minutes. That document would have taken you 2-3 days to write from scratch — if you ever got around to it."')

add_guardrail('Don\'t walk through more than 2 people. Everyone else shares in their pod or Slack channel.')

# ─── PHASE 6: CLOSE ───
add_phase_header(6, 'CLOSE', '1:10–1:20 — 10 min')

add_body('The Refinement Safety Net (2 min)')
add_body('Scroll to "When the Output Isn\'t Right" at the bottom of the Prompt Vault.')
add_script('"Quick note — if your output wasn\'t perfect, that\'s normal. The first output is a draft. Scroll to the bottom of the Prompt Vault page. There are six refinement prompts and the Feedback Formula. The pattern is simple: tell Claude what\'s working, what\'s not, and give it an example."')

add_body('Recap (2 min)')
add_script('"Here\'s what you built today:\n1. You now have a Claude account set up and ready\n2. You felt the difference between ChatGPT and Claude firsthand\n3. You have a Business Profile document that captures your entire business\n\nThat\'s three things you didn\'t have 90 minutes ago."')

add_body('Homework (1 min)')
add_script('"Here\'s what I need from you by [day]:\n1. Finish your Business Profile if you didn\'t complete it on the call\n2. Run Prompt 2 — the Ideal Client Profile\n3. Screenshot both documents and post them in your Slack channel\n\nThat\'s it. Two documents. Post proof."')

add_body('What\'s Next Teaser (1 min)')
add_body('Click to "What\'s Next" tab. Show the roadmap briefly.')
add_script('"Next session, we\'re taking these foundation documents and plugging them into a Claude Project — so every conversation starts with full context. Then we\'re building your content system. Come with an example of content you\'ve written that you\'re proud of."')

add_body('Round-Robin (4 min)')
add_script('"One word. How are you feeling about what we built today? Let\'s go around."')
add_body('Go person by person. React genuinely to each one.')

# ─── POST-SESSION ───
add_section_header('Post-Session — Send Within 24 Hours')
add_body('Paste this into the client Slack channel:')

post = doc.add_paragraph()
post.paragraph_format.left_indent = Cm(0.5)
post.space_after = Pt(4)
post_text = """Claude 101 — [Date]

What We Covered
1. Why Claude thinks differently from ChatGPT — diagnostic questions before answers
2. Hands-on: three raw prompts to feel the difference
3. Built your first Business Profile document using the Prompt Vault

What You Built
• Claude account (claude.ai)
• Business Profile document — your entire business mapped in one document

Your Homework
☐ Complete your Business Profile (if not finished on the call)
☐ Run Prompt 2 — Ideal Client Profile
☐ Screenshot both documents → post in your Slack channel
☐ Due by: [date]

Resources
• Full guide: claude.kaizencollective.com.au
• Recording: [link]

Next Session
[Date] — Setting up Claude Projects + building your content system.
Come with a piece of content you're proud of."""

run = post.add_run(post_text)
run.font.size = Pt(10)
run.font.color.rgb = CHARCOAL

# ─── TEACHING MOVES ───
add_section_header('Quick Reference: Teaching Moves')

table = doc.add_table(rows=7, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.LEFT

headers = ['When', 'Move', 'Example']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

moves = [
    ('After comparison', 'Cold-call echo', '"[Name], what did Claude do differently?"'),
    ('Before build sprints', 'Show the destination', 'Show a completed Business Profile first'),
    ('During build sprints', 'Walk the room', '"[Name], what question are you up to?"'),
    ('If energy drops', 'No-sevens rule', '"1-10, how confident? No sevens."'),
    ('When naming the AI gap', '"I\'m guilty of that"', '"I used ChatGPT the same way for months"'),
    ('Closing', 'One-word anchor', '"One word: how are you feeling?"'),
]

for i, (when, move, example) in enumerate(moves):
    row = table.rows[i + 1]
    row.cells[0].text = when
    row.cells[1].text = move
    row.cells[2].text = example
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

# ─── ANTI-PATTERNS ───
doc.add_paragraph()
add_section_header('Anti-Patterns to Watch')
add_checkbox('Don\'t spend more than 10 min on the Frame. Resist the urge to keep selling Claude.')
add_checkbox('Don\'t read every word on the page. Hit highlights, let them read later.')
add_checkbox('Don\'t walk through more than 3 people per build sprint.')
add_checkbox('Don\'t skip the recap. Say the three things they built.')
add_checkbox('Don\'t give vague homework. Specific action + specific deadline + post proof.')

# Save
output_path = '/Users/mariopaguio/Projects/claude-101-training/Claude-101-Session-Runsheet.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
